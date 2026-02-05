'''
pip install python-docx
'''
import subprocess
from docx import Document
from datetime import datetime
import os
from docx.shared import Pt  # 用于设置字体大小
from docx.oxml.ns import qn


kecheng='大数据集群搭建维护与数据处理'
cankaoziliao='《Hadoop系统搭建及项目实践》北京邮电大学出版社'

# kecheng='大数据数据库NoSQL数据库原理'
# cankaoziliao='《NoSQL数据库原理与应用案例教程》 航空工业出版社'

zhou = 18
file_path = 'E:\A\滨科教务\教学资料/4.教案/'+kecheng+'-（第'+str(zhou)+'周 4课时）教案.docx'
#----------------------------------------------------------------------------------------改为第几周
# 再生成详细教案内容，理论教学内容、实践教学环节、案例教学实施、教学重点突破策略、教学过程设计、教学效果评估，控制在1500字以内，不要格式，只输出内容

#课题
keti = '大数据技术体系总复习与综合能力强化'
#授课时间
shoukeshijian = '2025.12.31 5-8节'
#学情分析
xueqingfenxi='学生已经完成了从Hadoop集群搭建、HDFS原理与运维、MapReduce编程、Hive、Sqoop等组件学习到综合案例实践的全流程学习，对大数据技术生态有了较为全面的认识。但知识体系尚显零散，对不同技术的适用场景、优劣对比及组合运用缺乏系统化梳理。部分学生对前期内容有所遗忘，特别是在原理性知识和配置细节上。面对即将到来的课程考核，学生既需要巩固理论知识，也需要强化解决综合性问题的实践能力，并了解考核形式与重点。'
#教学内容
jiaoxueneirong = '''
Hadoop生态核心知识图谱梳理与各组件关系总览
HDFS核心机制、运维命令及常见故障排查要点回顾与强化
MapReduce编程模型深度复盘：从Shuffle过程到性能优化，典型编程模式（如过滤、聚合、连接）总结
Hive数据仓库工具的核心概念、架构、HQL常用操作及优化策略回顾
Sqoop数据迁移工具的使用场景与常用命令复盘
综合项目案例中技术选型与架构设计思路的再讨论与提炼
'''
#知识目标
zhishimubiao='''
系统掌握Hadoop生态的核心组件及其在大数据处理流程中的角色与协作关系。
巩固HDFS的架构原理、关键特性和运维管理知识。
深化理解MapReduce的工作机制、Shuffle过程及编程范式。
厘清Hive的数据模型、与传统数据库的异同及适用场景。
明确Sqoop在数据迁移中的定位和基本使用方法。
建立清晰的大数据项目技术选型与实施路径的知识框架。
'''
#能力目标
nenglimubiao='''
能够准确辨析不同大数据技术（如HDFS vs. 传统文件系统，MapReduce vs. Hive）的适用场景。
能够运用HDFS运维命令进行基本的管理和故障诊断。
能够独立设计与编写实现常见业务逻辑（如WordCount、去重、排序、连接）的MapReduce程序。
能够使用Hive SQL完成基本的数据查询、统计和分析任务。
能够根据一个简化的业务需求，勾勒出包含数据流和技术栈的解决方案草图。
具备应对课程理论考核与实践考核的基本解题与实操能力。
'''
#素质目标
suzhimubiao='''
培养对复杂技术体系的系统归纳与结构化思维能力。
提升在压力下（如备考）对已有知识进行有效提取和综合应用的心理素质。
强化严谨、细致的技术复习习惯，查漏补缺，追求对知识的精准理解。
树立诚信应考的纪律意识，强调扎实学习与真实能力的重要性。
'''
#教学重点
jiaoxuezhongdian='''
Hadoop生态整体知识结构与组件关联的梳理。
HDFS与MapReduce的核心原理与工作机制。
MapReduce典型编程模式与Hive核心操作的综合运用。
从需求到技术方案的设计思路提炼。
'''
#教学难点
jiaoxuenandian='''
帮助学生将零散的知识点整合成有机的、可迁移的知识体系。
针对学生的个性化知识薄弱点进行有效辅导和强化。
在有限时间内平衡知识串讲、习题解析和答疑等多种复习形式。
激发学生的主动复习动力，避免复习课变成教师的“独角戏”。
'''
#教学策略
jiaoxuecelv='采用“框架引领、问题驱动、分层互动”的复习策略。首先，教师利用思维导图展示完整的知识体系框架，建立全局认知。随后，以“问题链”（如“数据从MySQL到最终分析报告经历了什么？”）引导学生回顾关键技术和流程。设置“知识擂台”环节，学生分组竞答重点难点问题；开展“案例诊断”，对常见错误或低效方案进行集体剖析。提供分层练习题（基础、巩固、提高），满足不同层次学生需求。最后留出充足时间进行个性化答疑。'
#思政案例
sizhenganli='以“求真务实、诚信治学”为主题，结合期末复习与考核，引导学生认识到扎实掌握技术、诚实对待考试是对自己未来职业生涯负责的表现。大数据技术关乎数据真实与决策质量，从业者必须具备严谨求实的专业精神。鼓励学生通过认真复习巩固真才实学，拒绝投机取巧，将诚信意识内化为职业操守的基石。'
#教学反思
jiaoxuefansi='期末复习课信息密度大，需精心设计节奏，避免学生疲劳。思维导图等可视化工具对构建知识体系非常有效。互动环节（如擂台赛）能有效调动学生积极性，但需控制时间。分层练习和个性化答疑是关键，要关注后进生的需求。应强调复习方法（如抓住重点、理解原理而非死记命令），并提醒学生合理安排后续自主复习计划。教学效果很大程度上取决于学生课前自主梳理的程度，可提前布置梳理任务。'

#----------------------------------------------------------------------------------------
doc = Document('E:\A\滨科教务\教学资料/4.教案\教案模版.docx')

doc.paragraphs[1].text = '课程：'+kecheng+' 2025-2026学年第一学期 教师：卜凯凯'
paragraph = doc.paragraphs[1]
for run in paragraph.runs:
    run.bold = True

# 设置文档默认字体
style = doc.styles['Normal']
font = style.font

# 设置字体为仿宋_GB2312
font.name = '仿宋_GB2312'
font._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')

# 设置字体大小为小四号（12磅）
font.size = Pt(12)  # 小四号对应12磅

cell_content = doc.tables[0].rows[0].cells[1]
cell_content.text = keti.strip()  # 课题

cell_content = doc.tables[0].rows[1].cells[4]
cell_content.text = shoukeshijian.strip()  # 授课时间

cell_content = doc.tables[0].rows[2].cells[1]
cell_content.text = xueqingfenxi.strip()  # 学情分析
cell_content = doc.tables[0].rows[3].cells[1]
cell_content.text = jiaoxueneirong.strip()  # 教学内容

cell_content = doc.tables[0].rows[4].cells[2]
cell_content.text = zhishimubiao.strip()  # 知识目标
cell_content = doc.tables[0].rows[5].cells[2]
cell_content.text = nenglimubiao.strip()  # 能力目标
cell_content = doc.tables[0].rows[6].cells[2]
cell_content.text = suzhimubiao.strip()  # 素质目标

cell_content = doc.tables[0].rows[7].cells[1]
cell_content.text = jiaoxuezhongdian.strip()  # 教学重点
cell_content = doc.tables[0].rows[8].cells[1]
cell_content.text = jiaoxuenandian.strip()  # 教学难点
cell_content = doc.tables[0].rows[9].cells[1]
cell_content.text = jiaoxuecelv.strip()  # 教学策略
cell_content = doc.tables[0].rows[10].cells[1]
cell_content.text = sizhenganli.strip()  # 思政案例
cell_content = doc.tables[0].rows[11].cells[1]
cell_content.text = cankaoziliao.strip()  # 参考资料
cell_content = doc.tables[0].rows[12].cells[1]
cell_content.text = jiaoxuefansi.strip()  # 教学反思


        
# 保存修改
if os.path.exists(file_path):
    os.remove(file_path)
doc.save(file_path)
print("第"+str(zhou)+"周已生成")









