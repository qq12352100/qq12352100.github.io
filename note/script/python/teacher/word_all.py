import os
import sys
from docx import Document
from pathlib import Path

def read_docx_file(file_path):
    """读取.docx文件内容"""
    try:
        doc = Document(file_path)
        content = []
        
        # 提取段落内容
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # 只提取非空段落
                content.append(paragraph.text)
        
        # 提取表格内容
        for table in doc.tables:
            content.append("[表格开始]")
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                content.append(" | ".join(row_text))
            content.append("[表格结束]")
        
        return "\n".join(content)
    except Exception as e:
        return f"读取文件 {os.path.basename(file_path)} 时出错: {str(e)}"

def read_doc_file(file_path):
    """读取.doc文件内容（使用多种方法尝试）"""
    try:
        # 方法1: 尝试使用python-docx2txt库（如果可用）
        try:
            import docx2txt
            return docx2txt.process(file_path)
        except ImportError:
            pass
        
        # 方法2: 尝试使用win32com（仅限Windows，需要安装pywin32）
        try:
            import win32com.client
            
            # 创建Word应用对象
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False  # 不显示Word界面
            
            try:
                # 打开文档
                doc = word.Documents.Open(file_path)
                
                # 读取内容
                content = doc.Content.Text
                
                # 关闭文档
                doc.Close()
                
                return content
            finally:
                # 退出Word应用
                word.Quit()
        except ImportError:
            pass
        
        # 方法3: 尝试使用antiword（需要安装antiword软件）
        try:
            import subprocess
            
            # 检查系统是否安装了antiword
            result = subprocess.run(['antiword', '--version'], 
                                  capture_output=True, text=True)
            
            if result.returncode == 0:
                # 使用antiword提取文本
                result = subprocess.run(['antiword', file_path], 
                                      capture_output=True, text=True, encoding='utf-8')
                if result.returncode == 0:
                    return result.stdout
        except (FileNotFoundError, subprocess.SubprocessError):
            pass
        
        # 如果没有找到可用的方法
        return f"无法读取.doc文件 {os.path.basename(file_path)}，请安装以下之一：\n" \
               f"1. 安装docx2txt库: pip install docx2txt\n" \
               f"2. 安装pywin32库并确保已安装Microsoft Word: pip install pywin32\n" \
               f"3. 安装antiword工具并确保在PATH中\n" \
               f"或者将.doc文件另存为.docx格式"
    
    except Exception as e:
        return f"读取.doc文件 {os.path.basename(file_path)} 时出错: {str(e)}"

def read_word_files_to_txt(folder_path, output_file):
    """
    读取指定文件夹下的所有Word文件内容（包括.doc和.docx），并输出到一个txt文件中
    
    Args:
        folder_path (str): Word文件所在文件夹路径
        output_file (str): 输出的txt文件路径
    """
    try:
        # 检查文件夹是否存在
        if not os.path.exists(folder_path):
            print(f"文件夹 {folder_path} 不存在")
            return
        
        # 获取文件夹下所有Word文件
        word_extensions = ['.doc', '.docx']
        word_files = []
        
        for file in os.listdir(folder_path):
            file_path = os.path.join(folder_path, file)
            if os.path.isfile(file_path):
                file_ext = os.path.splitext(file)[1].lower()
                if file_ext in word_extensions:
                    word_files.append(file)
        
        if not word_files:
            print(f"在文件夹 {folder_path} 中未找到Word文件（支持.doc和.docx格式）")
            return
        
        print(f"找到 {len(word_files)} 个Word文件")
        
        # 按扩展名分类统计
        doc_count = sum(1 for f in word_files if f.lower().endswith('.doc'))
        docx_count = sum(1 for f in word_files if f.lower().endswith('.docx'))
        print(f"  - .doc文件: {doc_count} 个")
        print(f"  - .docx文件: {docx_count} 个")
        
        # 打开输出文件
        with open(output_file, 'w', encoding='utf-8') as txt_file:
            # 写入文件头
            txt_file.write(f"Word文件内容提取报告\n")
            txt_file.write(f"源文件夹: {folder_path}\n")
            txt_file.write(f"文件总数: {len(word_files)} 个\n")
            txt_file.write(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            txt_file.write("="*80 + "\n\n")
            
            # 遍历每个Word文件
            for i, word_file in enumerate(word_files, 1):
                word_file_path = os.path.join(folder_path, word_file)
                file_ext = os.path.splitext(word_file)[1].lower()
                
                print(f"正在处理: {word_file} ({i}/{len(word_files)})")
                
                # 写入文件名作为分隔
                txt_file.write(f"\n{'='*60}\n")
                txt_file.write(f"文件 {i}: {word_file} ({file_ext})\n")
                txt_file.write(f"{'='*60}\n\n")
                
                # 根据文件类型选择读取方法
                if file_ext == '.docx':
                    content = read_docx_file(word_file_path)
                else:  # .doc
                    content = read_doc_file(word_file_path)
                
                # 写入内容
                if content:
                    txt_file.write(content)
                else:
                    txt_file.write(f"文件 {word_file} 内容为空或读取失败\n")
                
                txt_file.write(f"\n{'='*60}\n\n")
        
        print(f"\n处理完成！内容已保存到: {output_file}")
        
        # 建议信息
        if doc_count > 0:
            print("\n注意：")
            print("由于.doc是旧格式文件，可能需要额外依赖才能读取。")
            print("如果遇到读取问题，请考虑将.doc文件转换为.docx格式。")
        
    except Exception as e:
        print(f"发生错误: {str(e)}")

def convert_doc_to_docx(folder_path):
    """
    可选功能：将文件夹中的.doc文件批量转换为.docx格式
    需要安装pywin32并确保已安装Microsoft Word
    """
    try:
        import win32com.client
        from win32com.client import constants
        
        # 获取所有.doc文件
        doc_files = [f for f in os.listdir(folder_path) 
                    if f.lower().endswith('.doc') and os.path.isfile(os.path.join(folder_path, f))]
        
        if not doc_files:
            print("未找到.doc文件")
            return
        
        print(f"找到 {len(doc_files)} 个.doc文件，开始转换...")
        
        # 创建Word应用对象
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False  # 不显示Word界面
        
        converted_count = 0
        
        for i, doc_file in enumerate(doc_files, 1):
            try:
                doc_path = os.path.join(folder_path, doc_file)
                docx_path = os.path.join(folder_path, os.path.splitext(doc_file)[0] + '.docx')
                
                print(f"转换: {doc_file} ({i}/{len(doc_files)})")
                
                # 打开文档
                doc = word.Documents.Open(doc_path)
                
                # 保存为.docx格式
                doc.SaveAs(docx_path, FileFormat=constants.wdFormatXMLDocument)
                
                # 关闭文档
                doc.Close()
                
                converted_count += 1
                print(f"  已保存为: {os.path.basename(docx_path)}")
                
            except Exception as e:
                print(f"  转换失败: {str(e)}")
        
        # 退出Word应用
        word.Quit()
        
        print(f"\n转换完成！成功转换 {converted_count}/{len(doc_files)} 个文件")
        
    except ImportError:
        print("未安装pywin32库，无法转换.doc文件")
        print("请安装: pip install pywin32")
    except Exception as e:
        print(f"转换过程中发生错误: {str(e)}")

if __name__ == "__main__":
    from datetime import datetime
    
    # 指定Word文件所在文件夹路径
    folder_path = r"E:\Temp\1"
    
    # 指定输出txt文件路径
    output_file = r"E:\Temp\word_contents.txt"
    
    # 执行转换
    read_word_files_to_txt(folder_path, output_file)
    
    # 可选：如果需要将.doc文件转换为.docx格式，可以取消下面的注释
    # convert_doc_to_docx(folder_path)