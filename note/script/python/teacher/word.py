'''
pip install python-docx
'''
import subprocess
from docx import Document
from datetime import datetime
import os
from docx.shared import Pt  # 用于设置字体大小
from docx.oxml.ns import qn


# 获取当前日期（本地时间）
today = datetime.today()
# 格式化为 "年-月-日" 中文格式
formatted_date = today.strftime("%Y年%m月%d日")

file_path = 'E:\A\管控项目资料\日报周报\工程日报-'+formatted_date+'.docx'

subprocess.run(["explorer", "E:\A\管控项目资料\日报周报"], shell=True)                           # 使用 explorer 命令打开文件夹

doc = Document('E:\A\管控项目资料\日报周报\日报.docx')

# 今日日报
today_content = '''
1、持续完善《系统管理操作手册》，补充典型操作示例及常见问题解答。
2、依据最新版汇报稿组织开展模拟演练，重点打磨逻辑结构、核心内容呈现及语言表达节奏。
3、针对分析指标汇总数据不一致时弹出的提示框，增加“取消”按钮以支持用户中止操作。
4、在情报采集模块新增功能：选择“包保责任人”时，默认限定为当前部门下的人员范围。
5、简化情报采集信息字段，移除用户职务信息的录入项。
6、优化平安态势报告导出功能，调整字体样式与表格排版。
7、针对系统运行过程中发现的其他问题，及时完成缺陷定位与修复。
8、持续推进督办中心相关功能开发与界面优化工作。
'''
# 明日计划
tomorrow_content = '''
1、继续完善汇报与培训相关材料及准备事项。
2、持续开展大模型相关参数的调优与验证工作，模型输出质量与响应效率。
3、进一步完善“三个中心”大屏的交互细节与视觉表现。
'''

ribao = '''


'''
# 读取所有段落
# print("=== 段落内容 ===")
# for para in doc.paragraphs:
    # print(para.text)

# 读取所有表格
# print("\n=== 表格内容 ===")
# for table in doc.tables:
    # for row in table.rows:
        # row_data = [cell.text for cell in row.cells]
        # print(row_data)
        
        
doc.paragraphs[1].text = '编制时间：'+formatted_date
# 设置字体为宋体，大小为小四（12磅）
for run in doc.paragraphs[1].runs:
    run.font.name = '宋体'
    run.font.size = Pt(12)
    # 解决中文显示问题的额外设置
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置今日日报内容
cell_today = doc.tables[0].rows[2].cells[1]
cell_today.text = today_content.strip()  # 先设置文本内容

# 格式化今日日报单元格（宋体小四）
for para in cell_today.paragraphs:
    for run in para.runs:
        run.font.name = '宋体'  # 西文字体
        run.font.size = Pt(12)  # 小四对应12磅
        # 设置中文字体（解决宋体显示问题）
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置明日计划内容
cell_tomorrow = doc.tables[0].rows[4].cells[1]
cell_tomorrow.text = tomorrow_content.strip()  # 先设置文本内容

# 格式化明日计划单元格（宋体小四）
for para in cell_tomorrow.paragraphs:
    for run in para.runs:
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        
# 保存修改
if os.path.exists(file_path):
    os.remove(file_path)
doc.save(file_path)











